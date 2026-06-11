"""Tools del dominio: documentos."""

from tools._core import *  # noqa: F401,F403

def list_documents(ocid: str, tool_context: ToolContext) -> dict:
    """Lista los documentos publicados en SEACE para esta convocatoria.

    Args:
        ocid: OCID de la convocatoria.

    Returns:
        Diccionario con documents (lista con id, title, documentType, url,
        format, has_b64).
    """
    cr = tool_context.state.get("ocds") or {}
    docs_b64 = tool_context.state.get("docs_b64") or {}
    docs = (cr.get("tender") or {}).get("documents") or []
    return {
        "n_documents": len(docs),
        "documents": [
            {
                "id": d.get("id"), "title": d.get("title"),
                "documentType": d.get("documentType"), "url": d.get("url"),
                "format": d.get("format"), "has_b64_preloaded": d.get("url") in docs_b64,
            }
            for d in docs[:5]
        ],
    }

def _analyze_pdf_layout(blob: bytes) -> dict:
    """Analiza la estructura de un PDF para detectar páginas cuyo contenido
    está rasterizado como imagen (es decir, el PDF tiene texto extraíble bajo,
    pero las páginas tienen imágenes grandes que cubren la mayoría del área).
    Estas páginas necesitan ser renderizadas a PNG y pasadas a Gemini Vision
    porque el OCR implícito del SDK sobre el bytestream del PDF muchas veces
    no recupera bien el contenido.

    Returns:
        {
          "n_pages": int,
          "needs_render_pages": [indices 0-based],
          "low_text_pages": [indices con <300 chars],
          "total_text_chars": int,
          "es_pdf_completamente_escaneado": bool,
        }
    """
    try:
        import fitz  # PyMuPDF
    except Exception as e:
        return {"error": f"pymupdf not available: {e}", "needs_render_pages": []}

    doc = fitz.open(stream=blob, filetype="pdf")
    needs_render: list[int] = []
    low_text: list[int] = []
    total_chars = 0
    n = len(doc)
    for i in range(n):
        page = doc[i]
        text = (page.get_text() or "").strip()
        n_chars = len(text)
        total_chars += n_chars
        if n_chars < 300:
            low_text.append(i)
        if n_chars >= 600:
            # página con texto suficiente — no requiere render
            continue
        # ¿hay imágenes cubriendo el área de la página?
        try:
            page_area = float(page.rect.width * page.rect.height) or 1.0
        except Exception:
            page_area = 1.0
        img_area = 0.0
        try:
            for img in page.get_images(full=True):
                xref = img[0]
                try:
                    for r in page.get_image_rects(xref):
                        img_area += float(r.width * r.height)
                except Exception:
                    # algunas builds no exponen image_rects con xref
                    pass
        except Exception:
            pass
        ratio = img_area / page_area if page_area else 0.0
        if ratio > 0.25:
            needs_render.append(i)
    es_escaneado = len(low_text) >= max(3, int(0.7 * n))
    doc.close()
    return {
        "n_pages": n,
        "needs_render_pages": needs_render,
        "low_text_pages": low_text,
        "total_text_chars": total_chars,
        "es_pdf_completamente_escaneado": es_escaneado,
    }

def _render_pdf_pages_to_png(
    blob: bytes, page_indices: list[int], dpi: int = 160,
) -> list[tuple[int, bytes]]:
    """Renderiza páginas específicas de un PDF a PNG bytes.

    Args:
        blob: bytes del PDF.
        page_indices: lista de índices 0-based de páginas a renderizar.
        dpi: resolución de render. 160 DPI = ~1300x1700 px en página A4, balance
             OCR vs tamaño.

    Returns:
        Lista de tuples (page_index_0based, png_bytes).
    """
    try:
        import fitz
    except Exception:
        return []
    doc = fitz.open(stream=blob, filetype="pdf")
    out: list[tuple[int, bytes]] = []
    try:
        for i in page_indices:
            if i < 0 or i >= len(doc):
                continue
            try:
                pix = doc[i].get_pixmap(dpi=dpi)
                out.append((i, pix.tobytes("png")))
            except Exception:
                continue
    finally:
        doc.close()
    return out

# ── Page-sharding: parte un PDF grande en sub-PDFs por rango de páginas ──
# Cada shard se procesa con su propia llamada Gemini (más chica, más rápida y
# sin riesgo de truncar el JSON), en paralelo (acotado por _throttle_gemini), y
# el merge de parse_document_pdf consolida items/firmantes/etc. por número.
PARSE_PAGES_PER_SHARD = int(os.getenv("PARSE_PAGES_PER_SHARD", "12"))
PARSE_SHARD_THRESHOLD = int(os.getenv("PARSE_SHARD_THRESHOLD", "16"))
PARSE_MAX_WORKERS = int(os.getenv("PARSE_MAX_WORKERS", "4"))
# Timeout por-llamada Gemini (ms): ninguna extracción de un shard puede colgarse
# más de esto. Visto en prod: una sola llamada de 7m36s congelaba toda la corrida.
PARSE_CALL_TIMEOUT_MS = int(os.getenv("PARSE_CALL_TIMEOUT_MS", "120000"))
# Techo POR-DOCUMENTO (s): una sola llamada a parse_document_pdf devuelve dentro
# de este presupuesto; los shards que no terminaron se marcan como timeout.
PARSE_OVERALL_TIMEOUT_S = int(os.getenv("PARSE_OVERALL_TIMEOUT_S", "600"))
# Techo GLOBAL del parser (s) across TODA la corrida: el primer parse_document_pdf
# fija un deadline compartido en state; cada documento respeta lo que queda. Así la
# SUMA de todos los PDFs no se come el wall de Cloud Run (1800s) y siempre se llega
# al writer + persist + force_flush. Debe dejar margen para el resto del pipeline.
PARSE_GLOBAL_BUDGET_S = int(os.getenv("PARSE_GLOBAL_BUDGET_S", "700"))


def _split_pdf_by_pages(blob: bytes, label: str,
                        pages_per_shard: int = PARSE_PAGES_PER_SHARD,
                        threshold: int = PARSE_SHARD_THRESHOLD,
                        overlap: int = 1) -> list[tuple[str, bytes]]:
    """Si el PDF supera `threshold` páginas, lo parte en sub-PDFs de
    ~`pages_per_shard` páginas (con `overlap` págs de solape para no cortar un
    ítem a la mitad). Devuelve [(label, blob)] tal cual si es chico o si falla."""
    try:
        import fitz
        src = fitz.open(stream=blob, filetype="pdf")
        n = src.page_count
        if n <= threshold:
            src.close()
            return [(label, blob)]
        shards: list[tuple[str, bytes]] = []
        for start in range(0, n, pages_per_shard):
            a = max(0, start - overlap)
            b = min(n - 1, start + pages_per_shard - 1)
            dst = fitz.open()
            dst.insert_pdf(src, from_page=a, to_page=b)
            shards.append((f"{label} [pp.{a + 1}-{b + 1}/{n}]", dst.tobytes()))
            dst.close()
            if b >= n - 1:
                break
        src.close()
        print(json.dumps({"pdf_sharded": label[:80], "n_pages": n, "n_shards": len(shards)}), flush=True)
        return shards or [(label, blob)]
    except Exception as e:
        print(json.dumps({"pdf_shard_error": str(e)[:160], "label": label[:80]}), flush=True)
        return [(label, blob)]


def _parse_single_pdf_with_gemini(blob: bytes, source_label: str) -> dict:
    """Procesa un PDF (bytes) con Gemini. Devuelve dict con extracción
    o {"error": ...}.

    Estrategia híbrida:
      1. Analiza la layout del PDF con PyMuPDF.
      2. Si hay páginas con texto < 300 chars + imagen cubriendo > 25% del área
         (típico cuando el REQUERIMIENTO está rasterizado como imagen embebida),
         renderiza esas páginas a PNG 160 DPI y se las pasa a Gemini como
         `image/png` parts adicionales junto con el PDF.
      3. Gemini hace OCR visual de alta calidad sobre los PNGs y combina con el
         texto extraíble del resto del PDF.
    """
    from google.genai import types as gtypes
    client = _gemini_client()

    # ── Análisis layout ─────────────────────────────────────────────
    layout = _analyze_pdf_layout(blob)

    # ── Document AI OCR (per-use): texto del shard en UNA llamada barata.
    #    Si funciona, a Gemini le pasamos el TEXTO (no PNGs rasterizados) →
    #    menos tokens, más rápido, mejor extracción. Si falla o no está
    #    configurado, caemos al render histórico de PNGs (sin romper nada). ──
    docai_text: str | None = None
    try:
        from tools.docai import extract_text_docai
        docai_text = extract_text_docai(blob)
    except Exception as _e:
        print(f"[docai] caller error ({type(_e).__name__}: {str(_e)[:160]})", flush=True)
        docai_text = None
    print(
        f"[docai] {source_label[:55]} → "
        + (f"TEXTO {len(docai_text):,} chars (Gemini sobre texto)" if docai_text
           else "None → fallback render+Gemini Vision"),
        flush=True,
    )

    # Con texto OCR no hace falta rasterizar páginas (el OCR ya cubre las
    # rasterizadas). Sin texto OCR → render como antes.
    pages_to_render = [] if docai_text else (layout.get("needs_render_pages") or [])[:_MAX_RENDER_PAGES]
    rendered: list[tuple[int, bytes]] = []
    if pages_to_render:
        rendered = _render_pdf_pages_to_png(blob, pages_to_render, dpi=160)
    schema = gtypes.Schema(
        type=gtypes.Type.OBJECT,
        properties={
            "cuantia_total": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True),
            "fuente_financiamiento": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
            "modalidad": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
            "tipo_documento_detectado": gtypes.Schema(
                type=gtypes.Type.STRING, nullable=True,
                description=(
                    "Tipo de documento OECE detectado a partir del contenido: "
                    "bases_administrativas, terminos_de_referencia, expediente_tecnico, "
                    "acta_buena_pro, contrato, propuesta_economica, otro."
                ),
            ),
            "contiene_requerimiento": gtypes.Schema(
                type=gtypes.Type.BOOLEAN, nullable=True,
                description=(
                    "True si en este PDF aparece la sección 'REQUERIMIENTO' / "
                    "'Términos de Referencia' / 'Especificaciones Técnicas' con detalle "
                    "técnico por ítem (marca, modelo, normas, dimensiones, certificaciones, "
                    "potencia, capacidad, materiales, etc.)."
                ),
            ),
            "items": gtypes.Schema(
                type=gtypes.Type.ARRAY,
                items=gtypes.Schema(
                    type=gtypes.Type.OBJECT,
                    properties={
                        "numero": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                            description=(
                                "Número del ítem como string: '1', '1.1', '1.2', '2', '2a'. "
                                "Si el OCDS tenía 1 ítem que agrupa varios productos, asignales "
                                "sub-numeración con punto (1.1, 1.2, ...)."
                            ),
                        ),
                        "padre_ocds_item": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                            description=(
                                "Si es desglose de un ítem padre del OCDS, número del padre. "
                                "Ej. canasta '1' con 10 sub-productos → cada sub tiene padre_ocds_item='1'."
                            ),
                        ),
                        "descripcion_corta": gtypes.Schema(
                            type=gtypes.Type.STRING,
                            description="TÍTULO del ítem (1 línea, ≤200 chars). NO meter specs acá.",
                        ),
                        "cantidad": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True),
                        "unidad": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                            description="UND, KG, M3, LITRO, SACO, BOLSA, etc."),
                        "precio_unitario_referencial": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True),
                        "cuantia_referencial_item": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True,
                            description="Cuantía total del ítem = precio_unitario × cantidad."),
                        "marca_o_modelo_exigido": gtypes.Schema(
                            type=gtypes.Type.STRING, nullable=True,
                            description=(
                                "Texto exacto de marca/modelo cuando aparece en el "
                                "documento, con sus eventuales modificadores ('o "
                                "similar', 'o equivalente'). Null si el ítem es "
                                "genérico. NO uses ejemplos de tu memoria — copiá "
                                "literal lo que aparece en el PDF."
                            ),
                        ),
                        "certificaciones_exigidas": gtypes.Schema(
                            type=gtypes.Type.ARRAY,
                            items=gtypes.Schema(type=gtypes.Type.STRING),
                            description=(
                                "Lista de certificaciones/normas exigidas al BIEN: 'EPA Tier 3', "
                                "'Homologación MTC', 'ROPS/FOPS', 'NTP 350.026', 'DIGESA', etc. "
                                "Cada entrada ≤80 chars (solo el nombre de la norma)."
                            ),
                        ),
                        "valores_tecnicos_clave": gtypes.Schema(
                            type=gtypes.Type.OBJECT,
                            nullable=True,
                            description=(
                                "Valores numéricos discretos extraídos del REQUERIMIENTO técnico. "
                                "Llená SOLO los que aparezcan; deja null el resto. Esto es lo que "
                                "market_price_agent usa para hacer queries específicas."
                            ),
                            properties={
                                "potencia_min_hp": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True),
                                "potencia_min_kw": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True),
                                "capacidad_volumen": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                                    description="Ej. '1.0 m3', '20 litros', '12L', '50 kg'."),
                                "capacidad_carga_ton": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True),
                                "peso_operativo_ton": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                                "alcance_m": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True),
                                "ano_fabricacion_min": gtypes.Schema(type=gtypes.Type.INTEGER, nullable=True),
                                "estado": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                                    description="'nueva sin uso', 'usada certificada', etc."),
                                "presentacion": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                                    description="Para consumibles: 'saco de 50kg', 'lata 140g', 'balde 20L'."),
                                "color": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                                "material": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                            },
                        ),
                        "garantia": gtypes.Schema(
                            type=gtypes.Type.OBJECT,
                            nullable=True,
                            properties={
                                "meses": gtypes.Schema(type=gtypes.Type.INTEGER, nullable=True),
                                "horas": gtypes.Schema(type=gtypes.Type.INTEGER, nullable=True,
                                    description="Para maquinaria: garantía expresada en horas de uso."),
                                "alcance": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                                    description="Ej. 'comercial', 'fábrica', 'integral con repuestos'."),
                            },
                        ),
                        "condiciones_entrega": gtypes.Schema(
                            type=gtypes.Type.OBJECT,
                            nullable=True,
                            properties={
                                "plazo_dias_calendario": gtypes.Schema(type=gtypes.Type.INTEGER, nullable=True),
                                "lugar_entrega": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                                "modalidad": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                                    description="'única', 'parcial', 'a demanda'."),
                            },
                        ),
                        "requisitos_postor": gtypes.Schema(
                            type=gtypes.Type.OBJECT,
                            nullable=True,
                            description="Requisitos al postor para participar (no al bien).",
                            properties={
                                "experiencia_minima_soles": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True),
                                "anos_experiencia_min": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True),
                                "n_contratos_similares": gtypes.Schema(type=gtypes.Type.INTEGER, nullable=True),
                                "certificaciones_postor": gtypes.Schema(
                                    type=gtypes.Type.ARRAY,
                                    items=gtypes.Schema(type=gtypes.Type.STRING),
                                    description="Ej. 'concesionario autorizado MTC', 'representante oficial de marca'."),
                                "infraestructura_exigida": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                                    description="Ej. 'taller propio con stock de repuestos', 'almacén refrigerado'."),
                                "personal_clave": gtypes.Schema(
                                    type=gtypes.Type.ARRAY,
                                    items=gtypes.Schema(type=gtypes.Type.STRING),
                                    description="Lista corta de personal exigido (ej. '1 mecánico CIP', '2 técnicos certificados')."),
                            },
                        ),
                        "penalidades": gtypes.Schema(
                            type=gtypes.Type.ARRAY,
                            items=gtypes.Schema(
                                type=gtypes.Type.OBJECT,
                                properties={
                                    "causal": gtypes.Schema(type=gtypes.Type.STRING),
                                    "monto_o_porcentaje": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                                    "base_calculo": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                                },
                                required=["causal"],
                            ),
                            description="Penalidades específicas mencionadas (mora, calidad, etc.).",
                        ),
                        "subitems": gtypes.Schema(
                            type=gtypes.Type.ARRAY,
                            description=(
                                "Si este ítem es un PAQUETE / LOTE / CANASTA que contiene N "
                                "productos físicos distintos, listalos acá con cantidad+unidad+"
                                "spec corta. Si el ítem es atómico (1 solo bien), dejá lista vacía."
                            ),
                            items=gtypes.Schema(
                                type=gtypes.Type.OBJECT,
                                properties={
                                    "descripcion": gtypes.Schema(type=gtypes.Type.STRING),
                                    "cantidad": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True),
                                    "unidad": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                                    "presentacion": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                                        description="Ej. 'saco 50kg', 'bolsa 1kg', 'lata 170g'."),
                                    "specs_clave": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                                        description="1 línea con specs (color, calidad, norma) si aparece."),
                                },
                                required=["descripcion"],
                            ),
                        ),
                        "requerimiento_tecnico_detallado": gtypes.Schema(
                            type=gtypes.Type.STRING, nullable=True,
                            description=(
                                "RESUMEN NARRATIVO denso (300-1000 chars) del requerimiento para "
                                "este ítem, integrando tipo de bien + marcas/modelos + valores "
                                "numéricos clave + certificaciones + garantía + requisitos del "
                                "postor + condiciones de entrega. Es lo que verá el writer del "
                                "dictamen. Los DETALLES estructurados ya van en los campos "
                                "discretos arriba — acá hacés la versión LEGIBLE. NO transcribas "
                                "literalmente 20K chars de specs por componente — comprimí lo "
                                "esencial. Si no hay requerimiento para este ítem, devolvé null."
                            ),
                        ),
                    },
                    required=["descripcion_corta"],
                ),
            ),
            "postores": gtypes.Schema(
                type=gtypes.Type.ARRAY,
                items=gtypes.Schema(
                    type=gtypes.Type.OBJECT,
                    properties={
                        "ruc": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                        "razon_social": gtypes.Schema(type=gtypes.Type.STRING),
                        "monto_oferta": gtypes.Schema(type=gtypes.Type.NUMBER, nullable=True),
                        "es_ganador": gtypes.Schema(type=gtypes.Type.BOOLEAN, nullable=True),
                        "item": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                    },
                    required=["razon_social"],
                ),
            ),
            "firmantes": gtypes.Schema(
                type=gtypes.Type.ARRAY,
                description=(
                    "Lista de personas que firman el documento. Aplica especialmente "
                    "a actas de buena pro, reportes de evaluación, contratos. "
                    "INCLUÍ a TODOS los firmantes que aparezcan al pie (presidente del "
                    "comité, miembros del comité, jefe de abastecimiento, gerente "
                    "general que aprueba, representante legal del proveedor)."
                ),
                items=gtypes.Schema(
                    type=gtypes.Type.OBJECT,
                    properties={
                        "nombre_completo": gtypes.Schema(type=gtypes.Type.STRING),
                        "dni": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                        "cargo": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                            description="Cargo institucional (ej. 'Jefe de Abastecimiento', 'Presidente del Comité')."),
                        "rol_en_documento": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                            description="Rol respecto al documento: 'aprobador', 'evaluador', 'presidente_comite', 'representante_proveedor', 'testigo'."),
                        "entidad": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                            description="Institución a la que pertenece (entidad contratante o empresa proveedora)."),
                        "fecha_firma": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                    },
                    required=["nombre_completo"],
                ),
            ),
            "comite_evaluacion": gtypes.Schema(
                type=gtypes.Type.ARRAY,
                description=(
                    "Composición formal del Comité de Selección o Comisión Evaluadora. "
                    "Si el comité no se detalla en el documento (común en contrataciones "
                    "directas), dejar lista vacía."
                ),
                items=gtypes.Schema(
                    type=gtypes.Type.OBJECT,
                    properties={
                        "nombre_completo": gtypes.Schema(type=gtypes.Type.STRING),
                        "cargo": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                        "rol": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                            description="presidente, miembro_titular, miembro_suplente, secretario."),
                        "certificacion_sican": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                    },
                    required=["nombre_completo"],
                ),
            ),
            "motivos_adjudicacion": gtypes.Schema(
                type=gtypes.Type.ARRAY,
                description=(
                    "Para CADA postor ganador, el motivo por el que ganó según lo "
                    "documentado en el acta o reporte de buena pro: criterio decisivo, "
                    "posición en ranking, observaciones de la evaluación, ajustes de "
                    "precio aplicados."
                ),
                items=gtypes.Schema(
                    type=gtypes.Type.OBJECT,
                    properties={
                        "ganador_razon_social": gtypes.Schema(type=gtypes.Type.STRING),
                        "ganador_ruc": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                        "item_adjudicado": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                        "criterio_decisivo": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                            description="ej. 'menor precio', 'único postor admitido', 'mejor calificación técnica', 'sorteo'."),
                        "posicion_ranking": gtypes.Schema(type=gtypes.Type.INTEGER, nullable=True),
                        "observaciones_evaluacion": gtypes.Schema(type=gtypes.Type.STRING, nullable=True,
                            description="Razones específicas que el comité documenta para esta adjudicación."),
                        "competidores_descalificados": gtypes.Schema(
                            type=gtypes.Type.ARRAY, items=gtypes.Schema(type=gtypes.Type.STRING),
                            description="Razones por las que se descalificaron otros postores (defectos en propuesta, no presentación, etc.)."),
                    },
                    required=["ganador_razon_social"],
                ),
            ),
            "lugar_fecha_acta": gtypes.Schema(
                type=gtypes.Type.OBJECT,
                nullable=True,
                description="Lugar y fecha de emisión del acta (cuando aplique).",
                properties={
                    "lugar": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                    "fecha": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                    "hora": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
                },
            ),
            "fundamento_legal": gtypes.Schema(
                type=gtypes.Type.ARRAY,
                items=gtypes.Schema(type=gtypes.Type.STRING),
            ),
            "resumen": gtypes.Schema(type=gtypes.Type.STRING, nullable=True),
        },
    )
    config = gtypes.GenerateContentConfig(
        temperature=0.0, top_p=0.1,
        response_mime_type="application/json",
        response_schema=schema, max_output_tokens=65535,  # Gemini 2.5 max
        http_options=gtypes.HttpOptions(timeout=PARSE_CALL_TIMEOUT_MS),  # techo por-llamada
        system_instruction=(
            "Sos un extractor experto en documentos del Sistema Electrónico de "
            "Contrataciones del Estado (SEACE) del Perú y del Organismo Especializado "
            "para las Contrataciones Eficientes del Estado (OECE / ex-OSCE). "
            "Procesás Bases Administrativas, Términos de Referencia (TDR), "
            "Especificaciones Técnicas (EETT), Expedientes Técnicos, Actas de Buena "
            "Pro, Contratos y Propuestas Económicas.\n"
            "\n"
            "ENTRADA: vas a recibir UN PDF en `application/pdf`. ADEMÁS, en muchos "
            "casos vas a recibir entre 1 y 30 imágenes PNG ADICIONALES. Esas imágenes "
            "son RENDERS A 160 DPI de páginas específicas del MISMO PDF cuyo "
            "contenido está rasterizado como imagen embebida (típico en bases del "
            "OECE — la sección REQUERIMIENTO viene casi siempre como imagen pegada en "
            "un PDF, NO como texto). El prompt del usuario te indica el número de "
            "página al que corresponde cada PNG. DEBÉS combinar el texto extraíble "
            "del PDF con el contenido de las imágenes para reconstruir la información "
            "completa. SIN las imágenes el REQUERIMIENTO no se ve.\n"
            "\n"
            "MISIÓN CRÍTICA: localizar la sección 'REQUERIMIENTO' (también llamada "
            "'Términos de Referencia', 'Especificaciones Técnicas', 'Características "
            "Técnicas del Bien', 'Características Técnicas del Servicio', 'Alcance "
            "del Servicio' o 'Características de la Obra') y EXTRAERLA A CAMPOS "
            "ESTRUCTURADOS — no transcribir, no perder información, no truncar.\n"
            "\n"
            "⚠ ANTI-ALUCINACIÓN — REGLA CRÍTICA DE INTEGRIDAD:\n"
            "  · NUNCA INVENTES contenido. Si no podés leer claramente el texto del\n"
            "    PDF (porque está mal rasterizado, las imágenes adjuntas no son\n"
            "    legibles, o el PDF aparenta estar dañado), respondé:\n"
            "      contiene_requerimiento=false\n"
            "      items=[]\n"
            "      resumen='No se pudo extraer información legible del documento'\n"
            "    NO completes con un caso 'genérico' o 'plantilla' (ej. servicio de\n"
            "    limpieza, kit de útiles, broca traumatológica) basado en tu memoria\n"
            "    de bases administrativas peruanas. Si NO está EN EL DOCUMENTO,\n"
            "    NO existe.\n"
            "  · El OBJETO DEL CONTRATO viene en el OCDS (entregado por la entidad)\n"
            "    y debe coincidir con lo que extraés del PDF. Si tu extracción\n"
            "    discrepa radicalmente del objeto OCDS (ej. OCDS dice 'codeína\n"
            "    fosfato' pero el PDF según vos habla de 'limpieza'), prioritariamente\n"
            "    revisá si te equivocaste leyendo el PDF — probablemente el PDF\n"
            "    SÍ habla de codeína y vos lo malinterpretaste.\n"
            "\n"
            "\n"
            "REGLAS DE EXTRACCIÓN ESTRUCTURADA (no perder NADA relevante):\n"
            "  · Sos un EXTRACTOR ESTRUCTURADO. Tu trabajo es leer el documento y\n"
            "    volcar TODA la información relevante a CAMPOS DISCRETOS. Cada dato\n"
            "    tiene su lugar específico en el schema:\n"
            "      - Marca/modelo →  `marca_o_modelo_exigido` (string)\n"
            "      - Normas técnicas →  `certificaciones_exigidas` (lista corta)\n"
            "      - Valores numéricos (potencia, capacidad, año, peso, alcance)\n"
            "        →  `valores_tecnicos_clave` (objeto con campos numéricos discretos)\n"
            "      - Garantía →  `garantia` (objeto: meses, horas, alcance)\n"
            "      - Plazo y lugar de entrega →  `condiciones_entrega` (objeto)\n"
            "      - Requisitos al postor (experiencia mínima en soles, años,\n"
            "        certificaciones del postor como 'concesionario MTC', infra)\n"
            "        →  `requisitos_postor` (objeto)\n"
            "      - Penalidades (causal + % + base de cálculo)\n"
            "        →  `penalidades` (lista de objetos)\n"
            "      - Si el ítem es un PAQUETE/LOTE/CANASTA con N productos\n"
            "        adentro (típico en bases de alimentos, kits escolares)\n"
            "        →  `subitems` (lista anidada, NO uses sub-numeración\n"
            "        en items[] para esto; usá esta lista)\n"
            "      - El narrativo legible que verá el dictamen\n"
            "        →  `requerimiento_tecnico_detallado` (resumen DENSO 300-1000 chars)\n"
            "\n"
            "  · NO transcribas literal. NO copies bloques de 20K chars. La spec\n"
            "    técnica viene en CAMPOS DISCRETOS — esos son tu output principal.\n"
            "    `requerimiento_tecnico_detallado` es solo el resumen para humanos.\n"
            "\n"
            "  · NO inventés especificaciones que no estén en el PDF/imágenes. Si un\n"
            "    campo no aparece, dejalo null.\n"
            "\n"
            "  · Asociás cada bloque de requerimiento con su NÚMERO DE ÍTEM "
            "(Ítem 1, Ítem 2, etc.). Si el documento tiene un solo ítem global, todo "
            "el REQUERIMIENTO se asocia a ese ítem.\n"
            "  · Si en el PDF figuran TABLAS de ítems (frecuente en bases para "
            "alimentos, uniformes, medicamentos, útiles), cada FILA de la tabla suele "
            "ser un ítem independiente — extraé un objeto en `items[]` por cada fila.\n"
            "  · NO inventás especificaciones que no estén en el PDF/imágenes.\n"
            "  · Si el documento NO contiene la sección REQUERIMIENTO (ej. es solo un "
            "acta o un contrato), dejá `contiene_requerimiento=false` y "
            "`requerimiento_tecnico_detallado=null` en cada ítem.\n"
            "  · Detectás marcas/modelos explícitos que aparezcan en el documento, "
            "y por separado las certificaciones/normas técnicas (MTC, Euro, Tier, "
            "ISO, NTP, DIGESA, SENASA, ASTM, EPA, etc.). Copiá los strings LITERALES "
            "del PDF — no traduzcas, no normalices, no completes con tu memoria.\n"
            "  · Identificás 'red flags' documentales: especificaciones que restringen "
            "competencia (marca única sin 'o similar', certificación atípica, plazos "
            "ultra-cortos, experiencia desproporcionada, lotes empaquetados sin "
            "justificación).\n"
            "\n"
            "EXTRACCIÓN OBLIGATORIA (CRÍTICO PARA EL PIPELINE INVESTIGATIVO):\n"
            "\n"
            "  · `firmantes`: TODA persona que firma el documento al pie. Suele estar\n"
            "    en la última página con título, nombre, cargo y firma. Capturá nombre\n"
            "    completo, DNI (si aparece), cargo institucional, rol respecto al\n"
            "    documento (aprobador / evaluador / presidente_comite / representante_proveedor /\n"
            "    testigo) y entidad. ESPECIALMENTE en actas de buena pro y contratos.\n"
            "    Este dato es lo que cruzaremos con el gerente del proveedor para\n"
            "    detectar parentezco o cargo previo compartido.\n"
            "\n"
            "  · `comite_evaluacion`: composición del Comité de Selección / Comisión\n"
            "    Evaluadora si el documento lo lista. Capturá nombre, cargo, rol\n"
            "    (presidente/miembro_titular/miembro_suplente/secretario) y certificación\n"
            "    SICAN si se menciona.\n"
            "\n"
            "  · `motivos_adjudicacion`: si es Acta de Buena Pro o Reporte, para CADA\n"
            "    postor ganador documentá: por qué ganó (criterio_decisivo: 'menor\n"
            "    precio', 'único postor admitido', 'mejor calificación técnica', 'sorteo'),\n"
            "    posición en ranking, observaciones del comité (descalificaciones de\n"
            "    otros postores, ajustes de precio, etc.), competidores_descalificados\n"
            "    con razones específicas.\n"
            "\n"
            "  · `lugar_fecha_acta`: lugar, fecha y hora de emisión cuando aplique.\n"
            "\n"
            "🚨 BASES ADMINISTRATIVAS ≠ ACTA DE BUENA PRO 🚨\n"
            "Las BASES ADMINISTRATIVAS se publican ANTES de la convocatoria. NO\n"
            "tienen ni firmantes del comité, ni motivos de adjudicación, ni acta\n"
            "de buena pro. Esos datos solo existen en documentos posteriores:\n"
            "Acta de Buena Pro, Contrato firmado, Cuadro de Evaluación.\n"
            "Si el documento que estás procesando es BASES ADMINISTRATIVAS o\n"
            "TÉRMINOS DE REFERENCIA, dejá `firmantes=[]`, `motivos_adjudicacion=[]`,\n"
            "`comite_evaluacion=[]`, `lugar_fecha_acta=null`. NO inventes un\n"
            "comité de selección, ni un Presidente del Comité, ni una fecha de\n"
            "firma — eso es alucinación.\n"
            "\n"
            "Si el documento NO es un acta / reporte / contrato (ej. son bases\n"
            "administrativas puras), dejá `firmantes=[]`, `motivos_adjudicacion=[]`,\n"
            "`comite_evaluacion=[]`, `lugar_fecha_acta=null`. NO inventes nombres,\n"
            "cargos, RUCs, entidades, lugares ni fechas. JAMÁS uses placeholders\n"
            "tipo 'LUGAR_ACTA_EXAMPLE', 'FIRMANTE_ACTA_EXAMPLE', 'CARGO_EXAMPLE',\n"
            "'POSTOR_1_EXAMPLE', 'ENTIDAD_CONTRATANTE_EXAMPLE', '12345678901',\n"
            "'Nombre Apellido', 'Funcionario X', 'Juan Perez Quispe', 'Juan Pérez',\n"
            "ni cualquier valor genérico — son alucinaciones. Si el dato no está\n"
            "en el PDF, el campo va vacío/null. PREFERÍ campo vacío a campo\n"
            "inventado.\n"
            "\n"
            "REGLA DURA — VALIDACIÓN DE FIRMANTE:\n"
            "Para emitir un objeto en `firmantes[]` DEBÉS tener AL MENOS UNO de:\n"
            "  (a) DNI explícito del firmante (8 dígitos visibles en el PDF), o\n"
            "  (b) Nombre de la entidad REAL del firmante (no 'Entidad\n"
            "      Contratante' literal, sino 'Municipalidad de X', 'Ministerio\n"
            "      de Y', con nombre concreto que aparece en el PDF), o\n"
            "  (c) Imagen/firma escaneada visible al pie del documento que\n"
            "      acompañe un nombre legible.\n"
            "Si NINGUNA de las tres se cumple, el firmante NO va al output.\n"
            "\n"
            "Devolvé SOLO JSON conforme al schema, sin markdown, sin fences."
        ),
    )

    # ── Armado de parts ─────────────────────────────────────────────
    docai_note = ""
    if docai_text:
        # CON Document AI: mandamos SOLO el texto OCR del documento completo
        # (sin el PDF ni PNGs) → 1 llamada Gemini lean. Cap 1M chars ≈ 250K
        # tokens (cubre ~330 págs; entra de sobra en el contexto de Gemini 2.5).
        parts: list = [gtypes.Part.from_text(text=(
            "═══ TEXTO OCR DEL DOCUMENTO COMPLETO (Google Document AI, alta fidelidad) ═══\n"
            + docai_text[:1000000]
        ))]
        docai_note = (
            "ARRIBA está el TEXTO OCR COMPLETO del documento entero, extraído por "
            "Google Document AI (incluye tablas y páginas que estaban rasterizadas "
            "como imagen). Es la ÚNICA fuente — extraé de ahí TODOS los ítems, "
            "especificaciones técnicas y banderas. Procesá el documento completo.\n\n"
        )
    else:
        # SIN Document AI (fallback): PDF + PNGs renderizados (Gemini Vision).
        parts = [gtypes.Part.from_bytes(data=blob, mime_type="application/pdf")]
    render_note = ""
    if rendered:
        pages_human = ", ".join(str(i + 1) for i, _ in rendered)
        render_note = (
            f"ADJUNTO {len(rendered)} imágenes PNG a 160 DPI correspondientes a las "
            f"páginas {pages_human} del PDF (en ese mismo orden). Esas páginas tienen "
            f"el contenido RASTERIZADO COMO IMAGEN dentro del PDF (no texto extraíble), "
            f"por eso te las paso por separado. Hacé OCR visual sobre ellas y "
            f"transcribí palabra por palabra cualquier especificación técnica, tabla "
            f"de ítems, listado de marcas, certificación, plazo o requisito que "
            f"contengan. Lo más probable es que el REQUERIMIENTO completo viva "
            f"acá.\n\n"
        )
        for idx, png in rendered:
            parts.append(gtypes.Part.from_bytes(data=png, mime_type="image/png"))

    layout_note = ""
    if layout and not layout.get("error"):
        layout_note = (
            f"Layout detectado por PyMuPDF: {layout['n_pages']} páginas, "
            f"{layout['total_text_chars']:,} chars de texto extraíble, "
            f"{len(layout.get('low_text_pages') or [])} páginas con < 300 chars "
            f"de texto, {len(layout.get('needs_render_pages') or [])} páginas "
            f"rasterizadas (con contenido en imagen).\n\n"
        )

    prompt = (
        f"PDF a procesar: {source_label}.\n\n"
        f"{layout_note}"
        f"{docai_note}"
        f"{render_note}"
        "Hacé esto, en este orden:\n"
        "\n"
        "PASO 1 — Identificá el tipo de documento (Bases Administrativas, TDR, EETT, "
        "Acta de Buena Pro, Contrato, Propuesta, etc.) y completá `tipo_documento_detectado`.\n"
        "\n"
        "PASO 2 — BUSCÁ la sección REQUERIMIENTO. Suele estar en el Capítulo III de las "
        "Bases Estándar del OECE, titulada 'REQUERIMIENTO' o 'TÉRMINOS DE REFERENCIA' o "
        "'ESPECIFICACIONES TÉCNICAS'. Si la encontrás (sea como texto del PDF o como "
        "contenido OCR de las imágenes adjuntas), marcá `contiene_requerimiento=true`.\n"
        "\n"
        "PASO 3 — Para CADA ítem del proceso (1, 2, 3... — si hay tabla de ítems en\n"
        "el documento o en las imágenes, cada fila es un ítem) extraé toda la info\n"
        "del REQUERIMIENTO técnico en CAMPOS DISCRETOS del schema:\n"
        "\n"
        "  IDENTIFICACIÓN:\n"
        "  · `numero` (string): '1', '1.1', '2', etc.\n"
        "  · `padre_ocds_item`: null si es ítem root del OCDS; el número del padre\n"
        "    si es un sub-ítem dentro de un ítem compuesto.\n"
        "  · `descripcion_corta` (≤200 chars): TÍTULO del ítem tal como aparece en\n"
        "    el documento (no inventes una más corta).\n"
        "  · `cantidad`, `unidad`, `precio_unitario_referencial`, `cuantia_referencial_item`.\n"
        "\n"
        "  MARCA Y NORMAS (copia LITERAL del documento, NO uses ejemplos de memoria):\n"
        "  · `marca_o_modelo_exigido` (string corto): texto exacto del documento si\n"
        "    el ítem exige una marca/modelo. 'sin marca' si el ítem es genérico.\n"
        "    Null si la sección no menciona requisito de marca.\n"
        "  · `certificaciones_exigidas` (lista corta): solo los códigos de norma\n"
        "    LITERALES del documento (cada string como aparece).\n"
        "\n"
        "  VALORES NUMÉRICOS CLAVE (objeto `valores_tecnicos_clave`):\n"
        "    Llená SOLO los campos que aparecen en el documento; resto null:\n"
        "      potencia_min_hp / potencia_min_kw / capacidad_volumen ('1.0 m3') /\n"
        "      capacidad_carga_ton / peso_operativo_ton / alcance_m /\n"
        "      ano_fabricacion_min / estado ('nueva sin uso') /\n"
        "      presentacion (para consumibles: 'saco 50kg', 'lata 140g') /\n"
        "      color / material.\n"
        "\n"
        "  GARANTÍA (objeto `garantia`):\n"
        "      meses (int) / horas (int — maquinaria) / alcance ('comercial', 'fábrica').\n"
        "\n"
        "  ENTREGA (objeto `condiciones_entrega`):\n"
        "      plazo_dias_calendario / lugar_entrega / modalidad ('única', 'parcial').\n"
        "\n"
        "  REQUISITOS AL POSTOR (objeto `requisitos_postor`):\n"
        "      experiencia_minima_soles / anos_experiencia_min /\n"
        "      n_contratos_similares / certificaciones_postor (ej. 'concesionario MTC',\n"
        "      'representante oficial de marca') / infraestructura_exigida /\n"
        "      personal_clave.\n"
        "\n"
        "  PENALIDADES (lista `penalidades`):\n"
        "      [{causal, monto_o_porcentaje, base_calculo}, ...].\n"
        "      Ej. {causal: 'mora en entrega', monto_o_porcentaje: '0.10%',\n"
        "           base_calculo: 'sobre monto del bien por cada día'}.\n"
        "\n"
        "  SUB-ÍTEMS — si el ítem es un PAQUETE/LOTE/CANASTA con varios productos:\n"
        "      `subitems`: [{descripcion, cantidad, unidad, presentacion, specs_clave},...]\n"
        "      Ej. para 'CANASTA DE ALIMENTOS':\n"
        "        [{descripcion: 'Arroz superior', cantidad: 2, unidad: 'BOLSA',\n"
        "          presentacion: '1kg', specs_clave: 'grano largo, taquillado'}, ...]\n"
        "\n"
        "  RESUMEN PARA EL DICTAMEN (`requerimiento_tecnico_detallado`, 300-1000 chars):\n"
        "    Narrativa legible que combina lo más relevante de los campos discretos\n"
        "    arriba, SIEMPRE construida sobre el contenido REAL del PDF que estás\n"
        "    procesando. Es para que el dictamen periodístico pueda citar texto.\n"
        "    NO es transcripción literal — es resumen denso.\n"
        "    🚨 Si el ítem es cemento, redactalo sobre cemento; si es uniforme, sobre\n"
        "    uniforme; si es servicio, sobre el servicio. JAMÁS arrastres ejemplos\n"
        "    de excavadora, maquinaria pesada, marcas, normas o magnitudes que no\n"
        "    aparezcan literalmente en el PDF procesado. Si el documento no tiene\n"
        "    suficiente detalle técnico, escribí menos — preferí 200 chars precisos\n"
        "    a 1000 chars inventados.\n"
        "\n"
        "  ⚠ NO TRUNQUES NI INVENTES. Si el documento tiene 80 páginas y describe\n"
        "    cilindros, voltajes, presiones específicas — esos van DENTRO de los\n"
        "    campos discretos correspondientes (potencia, certificaciones, etc.).\n"
        "    Pero NO necesitás meter todos los sub-bullets de 'Sistema Eléctrico',\n"
        "    'Sistema Hidráulico', 'Componentes Internos' en el resumen narrativo —\n"
        "    eso es ruido para el dictamen. Quedate con lo que define el producto.\n"
        "\n"
        "PASO 4 — Identificá `postores` (RUC, razón social, monto, ganador) si el PDF/imágenes los "
        "mencionan (típicamente en actas y propuestas).\n"
        "\n"
        "PASO 5 — `fundamento_legal` (lista de artículos citados textualmente por el\n"
        "documento: 'Art. 55.1.b Ley 32069', 'Art. 2 TUO Ley 30225', 'D.S. 009-2025-EF\n"
        "Art. 12', etc.). Solo LO QUE EL DOCUMENTO CITA — no interpretes si están bien\n"
        "invocados o no. Eso lo evalúa `document_legal_analyst_agent` aparte.\n"
        "\n"
        "PASO 6 — `cuantia_total`, `fuente_financiamiento`, `modalidad` (suma alzada / precios "
        "unitarios / esquema mixto / tarifas) y `resumen` (3-4 líneas).\n"
        "\n"
        "REGLAS FINALES:\n"
        "  · Sos EXTRACTOR puro: extraés HECHOS del documento. NO emitís juicios\n"
        "    legales ni banderas de riesgo — eso lo hace `document_legal_analyst_agent`\n"
        "    sobre tu output. Si tu extracción es buena (campos discretos completos,\n"
        "    frases textuales preservadas, sin invención), el analyst hace su trabajo\n"
        "    sin problema.\n"
        "  · Si el documento NO es una base / TDR / EETT, `contiene_requerimiento=false` y "
        "los requerimientos por ítem quedan en null.\n"
        "  · Devolvé SOLO JSON. SIN markdown, SIN fences, SIN texto antes ni después."
    )
    parts.append(gtypes.Part.from_text(text=prompt))

    try:
        with _throttle_gemini():
            resp = _gemini_call_with_retry(
                lambda: client.models.generate_content(
                    model=DEFAULT_GEMINI_MODEL, contents=parts, config=config,
                ),
            )
        text = (resp.text or "").strip()
        data = _safe_parse_json(text)
        if not isinstance(data, dict) or not data:
            return {"error": f"non-json response: {text[:200]}"}
        data["_size_bytes"] = len(blob)
        data["_source"] = source_label
        data["_pdf_layout"] = {
            "n_pages": layout.get("n_pages"),
            "rendered_pages_1based": [i + 1 for i, _ in rendered],
            "total_text_chars": layout.get("total_text_chars"),
            "es_pdf_completamente_escaneado": layout.get("es_pdf_completamente_escaneado"),
        }
        # LOGGING DE ORIGEN POR-PDF (caza-contaminación): qué documento produjo
        # qué ítems. Si un run de acelerómetros loguea 'CARNE DE RES', el source
        # apunta al PDF/URL cruzado → confirma fetch-chain vs parser.
        try:
            _it = data.get("items") or []
            _first = str((_it[0] if _it else {}).get("descripcion_corta") or "")[:90]
            print(f"[parser-src] src={source_label} · bytes={len(blob)} · "
                  f"n_items={len(_it)} · first='{_first}'", flush=True)
        except Exception:
            pass
        return data
    except Exception as e:
        return {"error": f"gemini failed: {str(e)[:150]}", "_source": source_label}

def _fetch_doc_bytes(document_url: str, tool_context: ToolContext) -> tuple[bytes | None, str, str | None]:
    """Obtiene los bytes de un documento intentando 4 caminos en orden:

      1. state['docs_b64'][url]      → b64 inline (path rápido, cuando el
                                       bridge pre-cargó por POST).
      2. state['doc_urls'][url]      → URL pública de GCS donde el bridge
                                       archivó el PDF. Descargamos con
                                       google-cloud-storage (auth automática
                                       con la SA del Cloud Run).
      3. OECE_RELAY_URL (env var)    → Cloudflare Worker que bypassa el WAF.
      4. Descarga directa            → último recurso (SEACE bloquea IPs GCP,
                                       devuelve 403 casi seguro).

    Returns:
        (bytes_o_None, fuente_str, mensaje_error_o_None)
        fuente_str ∈ {'inline_b64', 'gcs', 'relay', 'direct'}
    """
    # Helper: normalizar URL para matching robusto (sin query/fragment, sin trailing slash).
    # Necesario porque la URL del OCDS puede tener variantes (?v=1, /, encoding) que
    # difieren de la URL que el cliente registró en doc_urls/docs_b64.
    def _norm(u: str) -> str:
        try:
            from urllib.parse import urlparse, urlunparse
            p = urlparse(u)
            path = p.path.rstrip("/")
            return urlunparse((p.scheme.lower(), p.netloc.lower(), path, "", "", "")).lower()
        except Exception:
            return (u or "").lower()

    target_norm = _norm(document_url)

    def _lookup(d: dict):
        if document_url in d:
            return d[document_url]
        # Fallback: match por URL normalizada
        for k, v in d.items():
            if _norm(k) == target_norm:
                return v
        return None

    # 1) inline b64
    docs_b64 = tool_context.state.get("docs_b64") or {}
    pdf_b64 = _lookup(docs_b64)
    if pdf_b64:
        try:
            return base64.b64decode(pdf_b64, validate=True), "inline_b64", None
        except Exception as e:
            return None, "inline_b64", f"invalid_b64: {e}"

    # 2) GCS (preferido cuando no hay b64)
    doc_urls = tool_context.state.get("doc_urls") or {}
    gcs_target = _lookup(doc_urls)
    if gcs_target:
        blob, err = _download_from_gcs(gcs_target)
        if blob is not None:
            return blob, "gcs", None
        # si GCS falla, seguimos a relay/directo
        gcs_error = err
    else:
        gcs_error = None

    # 2.5) Downloader local (puente residencial peruano vía túnel) — el path
    #      CONFIABLE: SEACE bloquea IPs de datacenter (Cloud Run, colos de CF)
    #      con 403 pero acepta IPs residenciales PE. El servicio corre en la
    #      máquina del usuario, descarga con su IP, sube a GCS y devuelve gs://.
    dl_base = os.getenv("LOCAL_DOWNLOADER_URL", "").strip()
    if dl_base:
        try:
            ocid_hint = (
                tool_context.state.get("ocid")
                or tool_context.state.get("ocid_preloaded")
                or ""
            )
            r = requests.post(
                f"{dl_base.rstrip('/')}/download",
                json={"url": document_url, "ocid": ocid_hint},
                headers={"X-Vigia-Token": os.getenv("LOCAL_DOWNLOADER_TOKEN", "")},
                timeout=180,
            )
            if r.status_code == 200:
                gs = (r.json() or {}).get("gcs_path")
                if gs:
                    blob, err = _download_from_gcs(gs)
                    if blob is not None:
                        return blob, "local_downloader", None
                    downloader_error = f"downloader_gcs_read={err}"
                else:
                    downloader_error = "downloader_sin_gcs_path"
            else:
                downloader_error = f"downloader HTTP {r.status_code}"
        except Exception as e:
            downloader_error = f"downloader_exception: {str(e)[:120]}"
    else:
        downloader_error = None

    # 3) relay (Cloudflare Worker)
    relay_base = os.getenv("OECE_RELAY_URL", "").strip()
    if relay_base:
        try:
            relay_url = (
                f"{relay_base.rstrip('/')}/?url={requests.utils.quote(document_url, safe='')}"
            )
            r = requests.get(relay_url, headers=BROWSER, timeout=60)
            if r.status_code == 200 and len(r.content) > 100:
                return r.content, "relay", None
            relay_error = f"relay HTTP {r.status_code}"
        except Exception as e:
            relay_error = f"relay_exception: {str(e)[:120]}"
    else:
        relay_error = "OECE_RELAY_URL no configurado"

    # 4) directo
    try:
        r = requests.get(document_url, headers=BROWSER, timeout=30)
        if r.status_code == 200:
            return r.content, "direct", None
        direct_error = f"direct HTTP {r.status_code}"
    except Exception as e:
        direct_error = f"direct_exception: {str(e)[:120]}"

    # Todo falló — construimos error compuesto
    parts = []
    if gcs_target:
        parts.append(f"gcs={gcs_error or 'failed'}")
    if dl_base:
        parts.append(f"downloader={downloader_error}")
    parts.append(f"relay={relay_error}")
    parts.append(direct_error)
    return None, "failed", " · ".join(parts)

def _download_from_gcs(gcs_or_https_url: str) -> tuple[bytes | None, str | None]:
    """Descarga bytes desde GCS. Acepta dos formas de URL:
      · gs://bucket/path                          → usa google-cloud-storage SDK
      · https://storage.googleapis.com/bucket/p   → requests.get directo (más rápido,
        sin auth — funciona con buckets públicos de lectura que es nuestro caso).

    Para buckets privados con SA en Cloud Run, el path SDK se autentica con
    Application Default Credentials.
    """
    if gcs_or_https_url.startswith("https://"):
        try:
            r = requests.get(gcs_or_https_url, timeout=60)
            if r.status_code == 200 and len(r.content) > 100:
                return r.content, None
            return None, f"https GET HTTP {r.status_code}"
        except Exception as e:
            return None, f"https_exception: {str(e)[:200]}"

    # gs:// → SDK
    if not gcs_or_https_url.startswith("gs://"):
        return None, f"url GCS no parseable: {gcs_or_https_url[:120]}"
    try:
        from google.cloud import storage
    except Exception as e:
        return None, f"google-cloud-storage no instalado: {e}"
    rest = gcs_or_https_url[5:]
    if "/" not in rest:
        return None, f"url GCS no parseable: {gcs_or_https_url[:120]}"
    bucket_name, blob_path = rest.split("/", 1)
    try:
        client = storage.Client()
        bucket = client.bucket(bucket_name)
        blob = bucket.blob(blob_path)
        data = blob.download_as_bytes()
        if data and len(data) > 100:
            return data, None
        return None, "gcs blob vacío o muy chico"
    except Exception as e:
        return None, f"gcs_exception: {str(e)[:200]}"

def _is_docx_blob(blob: bytes) -> bool:
    """True si el blob es un archivo DOCX (Office Open XML).
    Un DOCX siempre empieza con PK (ZIP) y contiene `word/document.xml`.
    SEACE V3 publica algunas Bases Administrativas en DOCX en lugar de PDF.
    """
    if not blob or blob[:2] != b"PK":
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(blob)) as z:
            names = z.namelist()
            return any(n == "word/document.xml" for n in names)
    except zipfile.BadZipFile:
        return False

def _docx_to_synthetic_pdf(blob: bytes) -> bytes | None:
    """Convierte un DOCX en un PDF sintético procesable por el pipeline de
    Gemini. Estrategia híbrida:

      1. Extrae texto + tablas con python-docx → páginas de texto plano.
      2. Extrae imágenes embebidas (word/media/*) → páginas separadas con cada
         imagen renderizada full-page.

    El PDF sintético NO es visualmente bonito pero sí leíble por Gemini, que
    hará OCR Vision sobre las imágenes y leerá el texto plano directamente.

    Returns: bytes del PDF resultante o None si la conversión falla.
    """
    try:
        from docx import Document
        import fitz
    except Exception:
        return None

    # 1. Texto + tablas del DOCX
    text_chunks: list[str] = []
    try:
        d = Document(io.BytesIO(blob))
        for para in d.paragraphs:
            t = (para.text or "").strip()
            if t:
                text_chunks.append(t)
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line.strip(" |"):
                    text_chunks.append(line)
    except Exception:
        pass

    # 2. Imágenes embebidas (PNG/JPEG/etc en word/media/)
    images: list[tuple[str, bytes]] = []
    try:
        with zipfile.ZipFile(io.BytesIO(blob)) as z:
            for name in z.namelist():
                if name.startswith("word/media/") and not name.endswith("/"):
                    ext = name.rsplit(".", 1)[-1].lower()
                    if ext in ("png", "jpg", "jpeg", "gif", "bmp", "tif", "tiff"):
                        try:
                            images.append((name, z.read(name)))
                        except Exception:
                            continue
    except Exception:
        pass

    if not text_chunks and not images:
        return None

    # 3. Construir PDF sintético con PyMuPDF
    out_doc = fitz.open()
    full_text = "\n".join(text_chunks)
    # Páginas de texto: A4 portrait, 1700pt de alto, 595pt de ancho
    PAGE_W, PAGE_H = 595, 842
    MARGIN = 36
    FONT_SIZE = 9
    LINE_H = 12
    if full_text:
        # Dividir en chunks que caben en una página
        max_chars_per_page = 4500  # heurístico
        text_pages = [full_text[i:i + max_chars_per_page]
                      for i in range(0, len(full_text), max_chars_per_page)] or [""]
        for chunk in text_pages:
            page = out_doc.new_page(width=PAGE_W, height=PAGE_H)
            try:
                page.insert_textbox(
                    fitz.Rect(MARGIN, MARGIN, PAGE_W - MARGIN, PAGE_H - MARGIN),
                    chunk, fontsize=FONT_SIZE, fontname="helv",
                )
            except Exception:
                # Fallback: insertar como texto plano sin caja
                try:
                    page.insert_text((MARGIN, MARGIN + FONT_SIZE), chunk[:3000], fontsize=FONT_SIZE)
                except Exception:
                    pass
    # Páginas de imagen: una imagen por página (full-bleed)
    for name, img_bytes in images:
        try:
            page = out_doc.new_page(width=PAGE_W, height=PAGE_H)
            page.insert_image(
                fitz.Rect(MARGIN, MARGIN, PAGE_W - MARGIN, PAGE_H - MARGIN),
                stream=img_bytes,
            )
        except Exception:
            continue
    try:
        pdf_bytes = out_doc.tobytes()
    finally:
        out_doc.close()
    return pdf_bytes if pdf_bytes else None

def _images_to_synthetic_pdf(images: list[tuple[str, bytes]]) -> bytes | None:
    """Convierte una lista de imágenes (JPG/PNG/TIFF/BMP) en un PDF sintético
    de 1 imagen por página. Útil para bases SEACE que vienen como ZIP de
    escaneos sueltos. Gemini hace OCR Vision sobre cada página.
    """
    try:
        import fitz  # PyMuPDF
    except Exception:
        return None
    if not images:
        return None
    out = fitz.open()
    try:
        for name, img_bytes in images:
            try:
                # Insertar como página completa A4 (vertical o horizontal según aspect)
                pix = fitz.Pixmap(img_bytes)
                w, h = pix.width, pix.height
                # A4 = 595x842 pt; ajustamos orientación
                if w > h:
                    page = out.new_page(width=842, height=595)
                else:
                    page = out.new_page(width=595, height=842)
                rect = page.rect
                page.insert_image(rect, stream=img_bytes)
            except Exception:
                continue
        if len(out) == 0:
            return None
        pdf_bytes = out.tobytes()
    finally:
        out.close()
    return pdf_bytes if pdf_bytes else None

def parse_document_pdf(document_url: str, tool_context: ToolContext) -> dict:
    """Descarga un documento de SEACE y lo procesa con Gemini 2.5 Flash.
    SI el documento es un ZIP, descomprime y procesa TODOS los PDFs internos
    en paralelo (hasta 5 por archivo). Devuelve un consolidado.

    Estrategia de descarga (en orden):
      1. state['docs_b64'][url] — inline b64 del bridge (PDFs chicos).
      2. state['doc_urls'][url] — GCS bucket (PDFs grandes archivados por el bridge).
      3. OECE_RELAY_URL — Cloudflare Worker que bypassa el WAF.
      4. Directo — último recurso (suele dar 403 desde IPs de GCP).

    Args:
        document_url: URL del documento (`tender.documents[].url` del OCDS).

    Returns:
        Diccionario con `pdfs_procesados` (lista, uno por PDF interno), más
        consolidados: items_consolidados, postores_consolidados, red_flags,
        cuantia_total, fundamento_legal. Si la descarga falló, devuelve
        `error` + `_fetch_attempts` describiendo qué se intentó.
    """
    # ── Caché por URL (fix #2): si ya parseamos este documento en este run, NO
    #    re-descargamos ni re-OCR'eamos (Document AI + Gemini son caros). El primer
    #    parseo ya volcó su data a state['parser_raw_consolidated']; devolvemos el
    #    output compacto cacheado.
    _pdoc_cache = tool_context.state.get("_parsed_doc_cache") or {}
    if document_url in _pdoc_cache:
        print(f"[parse] cache HIT · ...{document_url[-44:]} — evito re-descarga/re-OCR", flush=True)
        return _pdoc_cache[document_url]

    blob, fetch_source, fetch_error = _fetch_doc_bytes(document_url, tool_context)
    if blob is None:
        return {
            "error": "download_failed",
            "url": document_url,
            "_fetch_attempts": fetch_error,
        }

    pdf_blobs: list[tuple[str, bytes]] = []
    docx_blobs_converted = 0  # contador para diagnóstico
    rar_blobs_extracted = 0

    # CASO 0 — RAR: extraer con rarfile + binario unrar-free (instalado en Dockerfile)
    if blob[:4] == b"Rar!":
        try:
            import rarfile
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".rar", delete=False) as tf:
                tf.write(blob)
                rar_path = tf.name
            try:
                with rarfile.RarFile(rar_path) as rf:
                    names = rf.namelist()
                    pdf_inside = [n for n in names if n.lower().endswith(".pdf")]
                    docx_inside = [n for n in names if n.lower().endswith(".docx")]
                    if not pdf_inside and not docx_inside:
                        return {
                            "error": "rar_no_pdfs_or_docxs",
                            "url": document_url,
                            "files_inside": names[:8],
                        }
                    for name in pdf_inside[:3]:
                        try:
                            pdf_blobs.append((f"{name} (de RAR)", rf.read(name)))
                            rar_blobs_extracted += 1
                        except Exception:
                            continue
                    for name in docx_inside[:3]:
                        try:
                            inner = rf.read(name)
                            synth = _docx_to_synthetic_pdf(inner)
                            if synth:
                                pdf_blobs.append((f"{name} (RAR→DOCX→PDF sintético)", synth))
                                rar_blobs_extracted += 1
                                docx_blobs_converted += 1
                        except Exception:
                            continue
            finally:
                try:
                    os.unlink(rar_path)
                except Exception:
                    pass
        except rarfile.BadRarFile:
            return {"error": "bad_rar", "url": document_url}
        except Exception as e:
            return {
                "error": f"rar_extraction_failed: {str(e)[:200]}",
                "url": document_url,
                "hint": "Verificá que unrar-free esté instalado en la imagen Docker.",
            }
        if not pdf_blobs:
            return {
                "error": "rar_no_extractable_content",
                "url": document_url,
                "hint": "El RAR no contenía PDFs ni DOCXs procesables.",
            }
        # OK, ya tenemos pdf_blobs poblados desde el RAR — saltar al
        # pipeline de procesamiento (línea ~3000+).
    # CASO 1 — el blob ES un DOCX (SEACE V3 publica algunas bases en DOCX)
    elif _is_docx_blob(blob):
        synth_pdf = _docx_to_synthetic_pdf(blob)
        if not synth_pdf:
            return {"error": "docx_conversion_failed", "url": document_url,
                    "hint": "python-docx no pudo extraer contenido del DOCX."}
        label = document_url.rsplit("/", 1)[-1][:80] + " (DOCX→PDF sintético)"
        pdf_blobs.append((label, synth_pdf))
        docx_blobs_converted += 1
    elif blob[:2] == b"PK":
        # CASO 2 — ZIP genérico: extraer PDFs, DOCXs, imágenes y ZIPs anidados.
        IMG_EXTS = (".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp")
        def _extract_from_zip(zip_bytes: bytes, prefix: str = "", depth: int = 0) -> None:
            """Extrae recursivamente PDFs/DOCXs/imágenes/ZIP-anidados.
            Modifica `pdf_blobs` y `docx_blobs_converted` por referencia.
            """
            nonlocal docx_blobs_converted
            if depth > 2:  # safety contra zip-bomb
                return
            try:
                with zipfile.ZipFile(io.BytesIO(zip_bytes)) as z:
                    names = z.namelist()
                    # PDFs directos
                    for name in [n for n in names if n.lower().endswith(".pdf")][:3]:
                        try: pdf_blobs.append((f"{prefix}{name}", z.read(name)))
                        except Exception: continue
                    # DOCXs → PDF sintético
                    for name in [n for n in names if n.lower().endswith(".docx")][:3]:
                        try:
                            inner = z.read(name)
                            synth = _docx_to_synthetic_pdf(inner)
                            if synth:
                                pdf_blobs.append((f"{prefix}{name} (DOCX→PDF)", synth))
                                docx_blobs_converted += 1
                        except Exception: continue
                    # ZIPs anidados — recursión hasta depth=2
                    for name in [n for n in names if n.lower().endswith(".zip")][:3]:
                        try:
                            inner_zip = z.read(name)
                            _extract_from_zip(inner_zip, prefix=f"{prefix}{name}/", depth=depth + 1)
                        except Exception: continue
                    # Imágenes → PDF sintético (1 imagen = 1 página). Útil para bases escaneadas.
                    image_blobs: list[tuple[str, bytes]] = []
                    for name in [n for n in names if n.lower().endswith(IMG_EXTS)][:10]:
                        try: image_blobs.append((name, z.read(name)))
                        except Exception: continue
                    if image_blobs:
                        synth = _images_to_synthetic_pdf(image_blobs)
                        if synth:
                            pdf_blobs.append((f"{prefix}{len(image_blobs)} imágenes (escaneo→PDF)", synth))
            except zipfile.BadZipFile:
                pass

        prev_len = len(pdf_blobs)
        _extract_from_zip(blob)
        if len(pdf_blobs) == prev_len:
            # Después de recursión total, no encontramos nada procesable
            try:
                with zipfile.ZipFile(io.BytesIO(blob)) as z:
                    files_inside = z.namelist()[:10]
            except Exception:
                files_inside = []
            return {"error": "zip_no_pdfs_docx_or_images",
                    "files_inside": files_inside,
                    "hint": "El ZIP no contenía PDFs, DOCXs ni imágenes procesables. Puede tener .doc legado, .xls, archivos firmados o estar corrupto."}
    elif blob[:4] == b"%PDF":
        label = document_url.rsplit("/", 1)[-1][:80]
        pdf_blobs.append((label, blob))
    else:
        return {"error": "not_a_pdf_or_zip_or_docx", "size": len(blob),
                "first_bytes_hex": blob[:8].hex()}

    # Estrategia de partición según el extractor disponible:
    #   · CON Document AI: NO shardeаmos para Gemini. El doc entero va como 1
    #     unidad → docai.extract_text_docai hace OCR (chunkeа a ≤30 págs INTERNO
    #     por el límite de la API) y CONCATENA el texto → Gemini estructura TODO
    #     en UNA sola llamada (90 págs ≈ 200-300K chars, entran de sobra).
    #   · SIN Document AI (fallback): page-sharding histórico — sub-PDFs por
    #     rango → varias llamadas Gemini Vision más chicas en paralelo.
    try:
        from tools.docai import docai_enabled
        _use_docai = docai_enabled()
    except Exception:
        _use_docai = False

    _expanded: list[tuple[str, bytes]] = []
    for (name, b) in pdf_blobs:
        if isinstance(b, (bytes, bytearray)) and b[:4] == b"%PDF" and not _use_docai:
            _expanded.extend(_split_pdf_by_pages(bytes(b), name))
        else:
            _expanded.append((name, b))  # con docai: doc entero (1 sola llamada Gemini)
    pdf_blobs = _expanded

    pdfs_procesados: list[dict] = [None] * len(pdf_blobs)
    # La concurrencia REAL la limita _throttle_gemini (semáforo global) para no
    # gatillar el rate limiter; el pool sólo encola los shards listos.
    #
    # TECHO GLOBAL: `as_completed(timeout=...)` corta a los PARSE_OVERALL_TIMEOUT_S
    # aunque algún shard se haya quedado lento. Antes esto usaba un `with` +
    # `fut.result(timeout=120)` que era código muerto (as_completed sólo entrega
    # futures YA terminados) y, peor, el `with` hacía shutdown(wait=True) → esperaba
    # igual a los hilos lentos. Resultado: una llamada Gemini de 7m+ colgaba toda la
    # corrida y nunca se llegaba al writer/persist/force_flush.
    # Presupuesto GLOBAL compartido: el 1er parse fija el deadline; los siguientes
    # respetan lo que queda. El techo efectivo de ESTE documento = min(techo
    # por-documento, presupuesto global restante), con un piso de 30s.
    _now = time.monotonic()
    _deadline = tool_context.state.get("_parse_deadline")
    if not isinstance(_deadline, (int, float)):
        _deadline = _now + PARSE_GLOBAL_BUDGET_S
        try:
            tool_context.state["_parse_deadline"] = _deadline
        except Exception:
            pass
    _eff_timeout = max(30.0, min(float(PARSE_OVERALL_TIMEOUT_S), _deadline - _now))

    ex = concurrent.futures.ThreadPoolExecutor(max_workers=PARSE_MAX_WORKERS)
    futures = {
        ex.submit(_parse_single_pdf_with_gemini, b, name): i
        for i, (name, b) in enumerate(pdf_blobs)
    }
    try:
        for fut in concurrent.futures.as_completed(futures, timeout=_eff_timeout):
            i = futures[fut]
            try:
                pdfs_procesados[i] = fut.result()
            except Exception as e:
                pdfs_procesados[i] = {"error": f"parallel exec failed: {str(e)[:120]}",
                                       "_source": pdf_blobs[i][0]}
    except concurrent.futures.TimeoutError:
        pendientes = [i for i in futures.values() if pdfs_procesados[i] is None]
        print(f"[parser] techo {_eff_timeout:.0f}s agotado (global restante {_deadline-_now:.0f}s) · "
              f"{len(pendientes)}/{len(futures)} shard(s) sin terminar → marcados timeout",
              flush=True)
        for i in pendientes:
            pdfs_procesados[i] = {"error": f"parse timeout (>{PARSE_OVERALL_TIMEOUT_S}s)",
                                   "_source": pdf_blobs[i][0]}
    finally:
        # No esperamos a los hilos lentos (no se pueden matar en Python); cancelamos
        # los encolados y seguimos el pipeline. Sus llamadas Gemini liberan el
        # semáforo de _throttle_gemini cuando terminen por su cuenta.
        ex.shutdown(wait=False, cancel_futures=True)

    def _item_key(it: dict):
        """Clave semántica para dedup de ítems (fix #1): descripción normalizada +
        cantidad. Evita que el MISMO ítem, numerado distinto en dos documentos
        ('2' vs '02', '1.0' vs '01'), sobreviva duplicado y duplique el trabajo del
        market agent. Devuelve None si no hay descripción ni número."""
        import unicodedata
        desc = (it.get("descripcion_corta") or it.get("descripcion") or "").strip().upper()
        desc = " ".join(desc.split())
        desc = "".join(c for c in unicodedata.normalize("NFKD", desc)
                       if not unicodedata.combining(c))
        if desc:
            return ("d", desc, it.get("cantidad"))
        num = it.get("numero")
        if num is not None and str(num).strip():
            return ("n", str(num).strip())
        return None

    items_all: list[dict] = []
    postores_all: list[dict] = []
    red_flags_all: list[str] = []
    fundamento_all: list[str] = []
    firmantes_all: list[dict] = []
    comite_all: list[dict] = []
    motivos_all: list[dict] = []
    lugar_fecha_acta = None
    cuantia_total = None
    algun_pdf_con_requerimiento = False
    for r in pdfs_procesados:
        if "error" in r:
            continue
        items_all.extend(r.get("items") or [])
        postores_all.extend(r.get("postores") or [])
        # red_flags_observadas: campo legacy, ya no se pide al parser. El análisis
        # legal lo hace `document_legal_analyst_agent` aparte. Si algún parser
        # legacy aún lo emite, lo recolectamos pero el flujo ya no depende de eso.
        red_flags_all.extend(r.get("red_flags_observadas") or [])
        fundamento_all.extend(r.get("fundamento_legal") or [])
        firmantes_all.extend(r.get("firmantes") or [])
        comite_all.extend(r.get("comite_evaluacion") or [])
        motivos_all.extend(r.get("motivos_adjudicacion") or [])
        if r.get("lugar_fecha_acta") and not lugar_fecha_acta:
            lugar_fecha_acta = r.get("lugar_fecha_acta")
        if r.get("contiene_requerimiento"):
            algun_pdf_con_requerimiento = True
        if cuantia_total is None and r.get("cuantia_total"):
            cuantia_total = r["cuantia_total"]

    # Items consolidados: dedup por clave SEMÁNTICA (descripción+cantidad), no por
    # número (fix #1). Priorizamos el `requerimiento_tecnico_detallado` más largo.
    consolidado_by_key: dict = {}
    for it in items_all:
        key = _item_key(it) or ("_unk", len(consolidado_by_key))
        actual = consolidado_by_key.get(key)
        if actual is None:
            consolidado_by_key[key] = dict(it)
            continue
        # Merge: campos no nulos del nuevo sobrescriben sólo si el actual no tiene
        for k, v in it.items():
            if v in (None, "", [], {}):
                continue
            cur = actual.get(k)
            # El requerimiento_tecnico_detallado más LARGO gana
            if k == "requerimiento_tecnico_detallado":
                if not cur or (isinstance(v, str) and len(v) > len(cur or "")):
                    actual[k] = v
            elif cur in (None, "", [], {}):
                actual[k] = v
    items_consolidados = list(consolidado_by_key.values())

    # Deduplicar firmantes por nombre+cargo
    seen_firm = set()
    firmantes_dedup = []
    for f in firmantes_all:
        if not isinstance(f, dict):
            continue
        # Usar `or ""` para tolerar valores None explícitos (que .get() con default
        # no captura — solo captura key-missing).
        key = ((f.get("nombre_completo") or "").strip().upper(),
               (f.get("cargo") or "").strip().upper())
        if key in seen_firm or not key[0]:
            continue
        seen_firm.add(key)
        firmantes_dedup.append(f)

    # Output COMPACTO para no inflar el context del orquestador. La data
    # completa va a state['parser_raw_consolidated'] (líneas abajo) y
    # build_market_input / persist_analysis_outputs la leen desde ahí.
    # Si el orquestador o el agent quieren ver detalle, leen state.
    output_dict = {
        "n_pdfs_procesados": len(pdfs_procesados),
        "n_pdfs_con_error": sum(1 for r in pdfs_procesados if "error" in r),
        "algun_pdf_con_requerimiento": algun_pdf_con_requerimiento,
        "n_items_consolidados": len(items_consolidados),
        "n_postores": len(postores_all),
        "n_firmantes": len(firmantes_dedup),
        "n_motivos_adjudicacion": len(motivos_all),
        "tiene_acta": bool(lugar_fecha_acta),
        "cuantia_total": cuantia_total,
        "_url": document_url,
        "_fetch_source": fetch_source,
        "_note": "Detalle completo en state['parser_raw_consolidated']",
    }
    # Detalle completo SOLO si hubo error en TODOS los PDFs (para debug).
    # Si todo ok, no devolvemos `pdfs_procesados` al caller.
    if output_dict["n_pdfs_con_error"] >= output_dict["n_pdfs_procesados"] and output_dict["n_pdfs_procesados"] > 0:
        output_dict["pdfs_procesados_debug"] = pdfs_procesados

    # GUARDAR el output ACUMULADO en state['parser_raw_consolidated'] para que
    # build_market_input y otros consumers puedan leer la data completa
    # SIN depender de que el agente document_parser la incluya íntegra en
    # su respuesta final (que se guarda en state['document_analysis']).
    # Cada vez que se procesa un PDF, mergeamos sus items/postores/firmantes
    # al acumulador.
    raw = tool_context.state.get("parser_raw_consolidated") or {
        "items_consolidados": [],
        "postores_consolidados": [],
        "firmantes_consolidados": [],
        "comite_evaluacion": [],
        "motivos_adjudicacion": [],
        "red_flags_observadas": [],
        "fundamento_legal": [],
        "documentos": [],
    }
    # Dedup items por clave SEMÁNTICA (fix #1) — antes era por `numero`, que dejaba
    # pasar el mismo ítem numerado distinto en dos documentos ('2' vs '02').
    existing_keys = {}
    for _it in raw["items_consolidados"]:
        _k = _item_key(_it)
        if _k is not None:
            existing_keys[_k] = _it
    for it in items_consolidados:
        k = _item_key(it)
        if k is None:
            raw["items_consolidados"].append(it)
            continue
        prev = existing_keys.get(k)
        if prev is None:
            raw["items_consolidados"].append(it)
            existing_keys[k] = it
        else:
            # Ya existe (mismo ítem desde otro doc): conservamos el requerimiento
            # técnico más largo y descartamos el duplicado.
            new_req = it.get("requerimiento_tecnico_detallado") or ""
            cur_req = prev.get("requerimiento_tecnico_detallado") or ""
            if len(new_req) > len(cur_req):
                prev["requerimiento_tecnico_detallado"] = new_req
    # Dedup postores por RUC
    seen_rucs = {p.get("ruc") for p in raw["postores_consolidados"] if p.get("ruc")}
    for p in postores_all:
        if p.get("ruc") and p["ruc"] not in seen_rucs:
            raw["postores_consolidados"].append(p)
            seen_rucs.add(p["ruc"])
        elif not p.get("ruc"):
            raw["postores_consolidados"].append(p)
    # Dedup firmantes por (nombre, cargo)
    seen_firm = {((f.get("nombre_completo") or "").upper(), (f.get("cargo") or "").upper())
                 for f in raw["firmantes_consolidados"]}
    for f in firmantes_dedup:
        key = ((f.get("nombre_completo") or "").upper(), (f.get("cargo") or "").upper())
        if key not in seen_firm:
            raw["firmantes_consolidados"].append(f)
            seen_firm.add(key)
    # Extends simples
    raw["comite_evaluacion"].extend(comite_all)
    raw["motivos_adjudicacion"].extend(motivos_all)
    raw["red_flags_observadas"] = list(dict.fromkeys(
        raw["red_flags_observadas"] + list(red_flags_all)
    ))
    raw["fundamento_legal"] = list(dict.fromkeys(
        raw["fundamento_legal"] + list(fundamento_all)
    ))
    raw["documentos"].append({
        "url": document_url,
        "n_pdfs_procesados": len(pdfs_procesados),
        "n_pdfs_con_error": sum(1 for r in pdfs_procesados if "error" in r),
        "algun_pdf_con_requerimiento": algun_pdf_con_requerimiento,
        "_fetch_source": fetch_source,
    })
    if lugar_fecha_acta and not raw.get("lugar_fecha_acta"):
        raw["lugar_fecha_acta"] = lugar_fecha_acta
    if cuantia_total and not raw.get("cuantia_total"):
        raw["cuantia_total"] = cuantia_total

    tool_context.state["parser_raw_consolidated"] = raw

    # Cachear el output compacto por URL (fix #2) para no re-parsear el mismo doc.
    _pdoc_cache[document_url] = output_dict
    tool_context.state["_parsed_doc_cache"] = _pdoc_cache

    return output_dict

# ── FunctionTool wrappers ──
list_documents_tool = FunctionTool(func=list_documents)
parse_document_pdf_tool = FunctionTool(func=parse_document_pdf)
